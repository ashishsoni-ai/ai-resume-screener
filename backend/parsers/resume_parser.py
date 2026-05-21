"""
Resume Parser — extracts raw text from PDF and DOCX files.

Why pdfplumber over PyPDF2?
- pdfplumber handles tables, multi-column layouts, and formatting much better.
- PyPDF2 often produces garbled text from modern resumes.
"""
import io
import logging
from pathlib import Path

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    full_text = "\n".join(text_parts).strip()
    if not full_text:
        raise ValueError("Could not extract text from PDF. It may be image-based (scanned).")
    return full_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    
    # Also extract text from tables (common in resume templates)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)
    
    full_text = "\n".join(paragraphs).strip()
    if not full_text:
        raise ValueError("Could not extract text from DOCX. File may be corrupted or empty.")
    return full_text


def extract_resume_text(file_bytes: bytes, filename: str) -> str:
    """
    Route to correct parser based on file extension.
    Returns cleaned resume text.
    """
    suffix = Path(filename).suffix.lower()
    
    if suffix == ".pdf":
        logger.info(f"Parsing PDF: {filename}")
        return extract_text_from_pdf(file_bytes)
    elif suffix in (".docx", ".doc"):
        logger.info(f"Parsing DOCX: {filename}")
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Upload PDF or DOCX.")
